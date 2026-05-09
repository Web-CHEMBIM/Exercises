"""
Extract all 98 MCQ questions from the Electric Circuits PDF and generate a DOCX
file with a 2-column table layout where each question is in its own bordered cell.
"""
import fitz  # PyMuPDF
import os
import io
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image

# -------------------------------------------------------------------
# Configuration
# -------------------------------------------------------------------
PDF_PATH = r'c:\Users\User\Desktop\Exercises\4.3 Electric circuits MCQ QP.pdf'
OUTPUT_DIR = r'c:\Users\User\Desktop\Exercises\question_images'
DOCX_PATH = r'c:\Users\User\Desktop\Exercises\4.3 Electric circuits MCQ.docx'

os.makedirs(OUTPUT_DIR, exist_ok=True)

# -------------------------------------------------------------------
# Step 1: Build question -> (page, image_index) mapping
# The text on each page contains the sequential question numbers.
# Images on each page correspond to those questions in order.
# -------------------------------------------------------------------

def build_question_map(doc):
    """
    Parse each page's text to find question numbers (integers that appear as
    standalone numbers in the text), then map each to the corresponding
    embedded image on that page.
    """
    question_map = {}  # q_num -> (page_idx, img_idx)
    
    # Skip page 0 (title page) and page 37 (Paper 2 header)
    skip_texts = ['Paper 1', 'Paper 2']
    
    for page_idx in range(doc.page_count):
        page = doc[page_idx]
        text = page.get_text()
        images = page.get_images(full=True)
        
        if not images:
            continue
        
        # Extract question numbers from text
        # Numbers in the text that are standalone (not part of paper codes)
        lines = text.strip().split('\n')
        q_nums = []
        for line in lines:
            line = line.strip()
            # Skip known non-question lines
            if any(skip in line for skip in [
                'PhysicsAndMathsTutor', 'Electric circuits', 
                '0625/', 'Paper', 'Questions are', 'extended only',
                'question', 'core and'
            ]):
                continue
            # Check if line is just a number
            if re.match(r'^\d+$', line):
                num = int(line)
                if 1 <= num <= 98:
                    q_nums.append(num)
        
        # Sort question numbers
        q_nums.sort()
        
        # Handle special cases where # images != # questions
        # Some pages have combined images for multiple questions
        if len(q_nums) == 0:
            continue
            
        if len(q_nums) == len(images):
            # Perfect 1:1 mapping
            for i, qn_val in enumerate(q_nums):
                question_map[qn_val] = (page_idx, i)
        elif len(q_nums) > len(images):
            # More questions than images - some questions share an image
            # In this case, typically 2 questions in 1 image block
            if len(images) == 1:
                # All questions on this page share one image
                for qn_val in q_nums:
                    question_map[qn_val] = (page_idx, 0)
            else:
                # Distribute questions across images as best as possible
                # Usually the first image has the first question(s) and so on
                imgs_per_q = len(images) / len(q_nums)
                for i, qn_val in enumerate(q_nums):
                    img_idx = min(int(i * imgs_per_q), len(images) - 1)
                    question_map[qn_val] = (page_idx, img_idx)
        else:
            # More images than questions - some images are sub-parts
            # Map questions to images sequentially
            for i, qn_val in enumerate(q_nums):
                if i < len(images):
                    question_map[qn_val] = (page_idx, i)
    
    return question_map


# -------------------------------------------------------------------
# Step 2: Extract question images from PDF
# -------------------------------------------------------------------

def extract_question_images(doc, question_map):
    """
    Extract each question's image from the PDF and save as PNG.
    Returns dict of q_num -> image_path.
    """
    image_paths = {}
    
    # Track which (page, img_idx) we've already extracted to avoid duplicates
    extracted = {}
    
    for q_num in sorted(question_map.keys()):
        page_idx, img_idx = question_map[q_num]
        key = (page_idx, img_idx)
        
        if key in extracted:
            # Multiple questions share same image
            image_paths[q_num] = extracted[key]
            continue
        
        page = doc[page_idx]
        images = page.get_images(full=True)
        
        if img_idx >= len(images):
            print(f"WARNING: Q{q_num} - img_idx {img_idx} out of range on page {page_idx+1}")
            continue
        
        xref = images[img_idx][0]
        base_image = doc.extract_image(xref)
        img_bytes = base_image["image"]
        ext = base_image["ext"]
        
        img_path = os.path.join(OUTPUT_DIR, f'Q{q_num}.{ext}')
        with open(img_path, 'wb') as f:
            f.write(img_bytes)
        
        image_paths[q_num] = img_path
        extracted[key] = img_path
    
    return image_paths


# -------------------------------------------------------------------
# Step 3: Handle pages where questions share images
# For pages where 2+ questions are in one image, we need to render
# the page and crop individual question regions.
# -------------------------------------------------------------------

def extract_shared_questions_by_rendering(doc, question_map, image_paths):
    """
    For questions that share an image, render the full page and crop
    individual question regions based on image bounding boxes.
    """
    # Find questions that share images
    shared_groups = {}
    for q_num, (page_idx, img_idx) in question_map.items():
        key = (page_idx, img_idx)
        if key not in shared_groups:
            shared_groups[key] = []
        shared_groups[key].append(q_num)
    
    # For groups with multiple questions sharing one image,
    # render the page and try to split
    for (page_idx, img_idx), q_nums in shared_groups.items():
        if len(q_nums) <= 1:
            continue
        
        q_nums.sort()
        print(f"Shared image on page {page_idx+1}: Questions {q_nums}")
        
        # For shared images, we'll keep the full image for all questions
        # since it's hard to auto-split correctly.
        # The image already contains all the questions.


# -------------------------------------------------------------------
# Step 4: Alternative approach - Render full pages and crop questions
# based on text positions
# -------------------------------------------------------------------

def render_and_crop_all(doc):
    """
    Render each page at high DPI and crop individual question regions.
    This gives cleaner results than extracting embedded images.
    Uses the text positions of question numbers to determine crop regions.
    """
    image_paths = {}
    DPI = 200
    
    for page_idx in range(doc.page_count):
        page = doc[page_idx]
        text = page.get_text()
        images = page.get_images(full=True)
        
        if not images:
            continue
        
        # Get question numbers on this page
        lines = text.strip().split('\n')
        q_nums = []
        for line in lines:
            line = line.strip()
            if any(skip in line for skip in [
                'PhysicsAndMathsTutor', 'Electric circuits', 
                '0625/', 'Paper', 'Questions are', 'extended only',
                'question', 'core and'
            ]):
                continue
            if re.match(r'^\d+$', line):
                num = int(line)
                if 1 <= num <= 98:
                    q_nums.append(num)
        
        q_nums.sort()
        if not q_nums:
            continue
        
        # Get image bounding boxes
        img_info = page.get_image_info()
        
        if len(q_nums) == len(img_info):
            # 1:1 mapping - extract each embedded image directly
            for i, q_num in enumerate(q_nums):
                xref = images[i][0]
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                ext = base_image["ext"]
                img_path = os.path.join(OUTPUT_DIR, f'Q{q_num}.{ext}')
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                image_paths[q_num] = img_path
        elif len(q_nums) > len(img_info):
            # More questions than images - some questions share an image
            # Render the page and use the full image for each
            if len(img_info) == 1:
                # Single image contains all questions
                xref = images[0][0]
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                ext = base_image["ext"]
                img_path = os.path.join(OUTPUT_DIR, f'Q{"_".join(str(q) for q in q_nums)}.{ext}')
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                for q_num in q_nums:
                    image_paths[q_num] = img_path
            else:
                # Distribute: render page and crop
                pix = page.get_pixmap(dpi=DPI)
                page_img = Image.open(io.BytesIO(pix.tobytes("png")))
                scale = DPI / 72.0  # PDF coords are in 72 DPI
                
                # Assign questions to closest images
                for i, q_num in enumerate(q_nums):
                    # Find the closest image bbox
                    img_idx = min(i, len(img_info) - 1)
                    bbox = img_info[img_idx]['bbox']
                    
                    # Scale bbox to pixel coords
                    x0 = int(bbox[0] * scale) - 10
                    y0 = int(bbox[1] * scale) - 10
                    x1 = int(bbox[2] * scale) + 10
                    y1 = int(bbox[3] * scale) + 10
                    
                    # Clamp
                    x0 = max(0, x0)
                    y0 = max(0, y0)
                    x1 = min(page_img.width, x1)
                    y1 = min(page_img.height, y1)
                    
                    cropped = page_img.crop((x0, y0, x1, y1))
                    img_path = os.path.join(OUTPUT_DIR, f'Q{q_num}.png')
                    cropped.save(img_path)
                    image_paths[q_num] = img_path
        else:
            # More images than questions
            # Take only the first len(q_nums) images
            for i, q_num in enumerate(q_nums):
                if i < len(images):
                    xref = images[i][0]
                    base_image = doc.extract_image(xref)
                    img_bytes = base_image["image"]
                    ext = base_image["ext"]
                    img_path = os.path.join(OUTPUT_DIR, f'Q{q_num}.{ext}')
                    with open(img_path, 'wb') as f:
                        f.write(img_bytes)
                    image_paths[q_num] = img_path
    
    return image_paths


# -------------------------------------------------------------------
# Step 5: Create the DOCX with 2-column table layout
# -------------------------------------------------------------------

def set_cell_border(cell, **kwargs):
    """
    Set cell border on a table cell.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs["val"]}" '
            f'w:sz="{attrs["sz"]}" w:space="0" w:color="{attrs["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def create_docx(image_paths, total_questions=98):
    """
    Create a Word document with 2-column table layout.
    Each question is in a bordered cell with its image.
    """
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run('Exercise xx')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    
    # Calculate available width for each column in inches
    page_width_emu = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    col_width_inches = page_width_emu / 914400 / 2  # Convert EMU to inches, divide by 2 columns
    
    # Create questions in pairs (2 per row)
    q_nums = sorted(image_paths.keys())
    
    # Pair up questions
    pairs = []
    for i in range(0, len(q_nums), 2):
        if i + 1 < len(q_nums):
            pairs.append((q_nums[i], q_nums[i + 1]))
        else:
            pairs.append((q_nums[i], None))
    
    for pair_idx, (q1, q2) in enumerate(pairs):
        # Create a 1-row, 2-column table for each pair
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table width
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        
        # Set borders on the table
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
        
        # Process each cell
        for col_idx, q_num in enumerate([q1, q2]):
            cell = table.cell(0, col_idx)
            
            if q_num is None:
                continue
            
            # Add question number label
            p = cell.paragraphs[0]
            run = p.add_run(f'1')
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            
            # Add small space
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            
            # Add question number in superscript-like format
            # Actually, let's use the marks indicator
            mark_run = p.add_run(f'¤')
            mark_run.font.size = Pt(7)
            
            # Add the question image
            if q_num in image_paths and os.path.exists(image_paths[q_num]):
                img_path = image_paths[q_num]
                
                # Get image dimensions to scale properly
                with Image.open(img_path) as img:
                    img_w, img_h = img.size
                
                # Calculate scaling to fit in cell
                # Cell width is approximately half the page width minus padding
                max_width_inches = col_width_inches - 0.3  # subtract padding
                max_width_inches = min(max_width_inches, 3.5)
                
                # Scale proportionally
                aspect = img_h / img_w
                width = Inches(max_width_inches)
                
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                p2.paragraph_format.space_before = Pt(2)
                run2 = p2.add_run()
                run2.add_picture(img_path, width=width)
            
            # Set cell padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="60" w:type="dxa"/>'
                f'<w:left w:w="60" w:type="dxa"/>'
                f'<w:bottom w:w="60" w:type="dxa"/>'
                f'<w:right w:w="60" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(tcMar)
        
        # Add small spacing between tables
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(1)
        spacer.paragraph_format.space_after = Pt(1)
        pf = spacer.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
    
    doc.save(DOCX_PATH)
    print(f'Saved DOCX to: {DOCX_PATH}')


# -------------------------------------------------------------------
# Main
# -------------------------------------------------------------------

def main():
    print("Opening PDF...")
    doc = fitz.open(PDF_PATH)
    print(f"PDF has {doc.page_count} pages")
    
    print("\nExtracting question images...")
    image_paths = render_and_crop_all(doc)
    
    print(f"\nExtracted {len(image_paths)} question images")
    
    # Check for missing questions
    all_q = set(range(1, 99))
    found_q = set(image_paths.keys())
    missing = all_q - found_q
    if missing:
        print(f"Missing questions: {sorted(missing)}")
    
    print("\nCreating DOCX...")
    create_docx(image_paths)
    
    doc.close()
    print("Done!")


if __name__ == '__main__':
    main()
